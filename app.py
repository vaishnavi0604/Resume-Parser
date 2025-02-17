from flask import Flask, request, jsonify, render_template_string
import os
import re
import PyPDF2
import pdfplumber
from docx import Document
import spacy
import pandas as pd

nlp = spacy.load('en_core_web_sm')

app = Flask(__name__)

UPLOAD_FOLDER = 'uploads/'
ALLOWED_EXTENSIONS = {'pdf', 'docx', 'txt'}
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_pdf(pdf_path):
    text = ""
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            text += page.extract_text() + "\n" if page.extract_text() else ""

    return text.strip()

def extract_text_from_docx(docx_path):
    doc = Document(docx_path)
    text = []

    # Extract text from paragraphs
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text.append(paragraph.text.strip())

    # Extract text from tables (DOB might be inside a table)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text.append(cell.text.strip())

    return "\n".join(text)  # Join all text with newlines

def extract_text_from_txt(txt_path):
    with open(txt_path, 'r') as file:
        return file.read()

def extract_name(text):
    # Improved name extraction that handles both formats
    name_label_match = re.search(
        r'(?i)(?:Name|^)\s*:*\s*([A-Z][a-z]+\s+[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)\b', 
        text
    )
    
    if name_label_match:
        full_name = name_label_match.group(1).strip()
    else:
        # Fallback for bold headers in DOCX
        first_line = text.split('\n')[0].replace('**', '').strip()
        name_pattern = r'^([A-Z][a-z]+\s+[A-Z][a-z]+)'
        first_line_match = re.match(name_pattern, first_line)
        full_name = first_line_match.group(1).strip() if first_line_match else "Not provided"

    parts = full_name.split()
    return (parts[0], ' '.join(parts[1:])) if len(parts) >= 2 else (full_name, "Not provided")

def extract_email(text):
    email_pattern = r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}"
    return re.findall(email_pattern, text)

def extract_skills(text):
    # Updated to handle different section headers and more skills
    skills_section = re.search(r'(?i)(?:Technical Skills|Skills? & Abilities|Skill Set)[\s:-]*(.*?)(?=\n\s*(?:Experience|Education|Projects|$))', text, re.DOTALL)
    if skills_section:
        return re.findall(r'\b(HTML|CSS|JavaScript|Bootstrap|C#|ASP\.NET|ADO\.NET|MVC|SQL Server|Entity Framework|LINQ)\b', skills_section.group(1), re.IGNORECASE)
    return []   

def extract_phone_number(text):
    # Improved regex to match phone numbers with headings like "Contact Info", "Phone", "Mobile", etc.
    phone_pattern = r"(?i)(?:(?:Contact(?:\s*Info)?|Phone|Mobile|Tel|Cell)[\s:.-]*)?(\+?\d{1,3}[-.\s]?\(?\d{2,4}\)?[-.\s]?\d{3,4}[-.\s]?\d{4})"

    matches = re.findall(phone_pattern, text)

    # Filter out empty results and normalize numbers
    phone_numbers = [match.strip() for match in matches if match.strip()]
    
    return phone_numbers if phone_numbers else ["Not provided"]

def extract_address(text):
    # Improved address extraction that skips name line
    address_pattern = r'(?i)(?:Address|Correspondence Address)[:\s-]*([^\n]+)(?:\n\s*([^\n]+))?'
    match = re.search(address_pattern, text)
    
    if match:
        return ' '.join([g.strip() for g in match.groups() if g])
    
    # Fallback for resumes without address label
    address_candidates = [
        line for line in text.split('\n') 
        if re.search(r'\d{6}', line) and not re.search(r'(name|email|phone)', line.lower())
    ]
    return address_candidates[0] if address_candidates else "Not provided"

def extract_dob(text):
    dob_patterns = [
        
       r'(?i)\b(?:dob|date\s*of\s*birth|birth\s*date|d\.o\.b)\s*[:\-\s]*\s*(\d{1,2}[-\/\s.]+\d{1,2}[-\/\s.]+\d{4})\b',
        r'\b(\d{1,2}[-\/\s.]\d{1,2}[-\/\s.]\d{4})\b',
        r'\b(\d{4}[-\/\s.]\d{1,2}[-\/\s.]\d{1,2})\b'
    ]

    for pattern in dob_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            return match.group(1).strip()  # Use group(1) to capture the date part

    return "Not provided"
    
def extract_father_name(text):
    # Strict regex that stops at line end or next label
    father_pattern = r'(?i)father[\'’s]*\s*name[:\s-]+\s*([^\n]+?)(?=\s*\n|$)'
    match = re.search(father_pattern, text)
    return match.group(1).strip() if match else "Not provided"

def extract_gender(text):
    gender_pattern = r'(?i)gender[:\s-]+([A-Za-z]+)'
    match = re.search(gender_pattern, text)
    return match.group(1).strip() if match else "Not provided"

def extract_languages(text):
    lang_pattern = r'(?i)languages?\s*(?:known|proficiency)[:\s-]+([A-Za-z,\s]+?)(?=\n|$)'
    match = re.search(lang_pattern, text)
    if match:
        return [lang.strip() for lang in re.split('[,/]', match.group(1))]
    return ["Not provided"]

def extract_certifications(text):
    cert_pattern = r'(?i)(?:Certifications?|Certificates?)[:\s-]*(.*?)(?=\n\s*(?:Experience|Education|$))'
    cert_section = re.search(cert_pattern, text, re.DOTALL)
    if cert_section:
        return re.findall(r'- (.*?)(?=\n|$)', cert_section.group(1))
    return []

def extract_profiles(text):
    profiles = {
        'LinkedIn' : re.findall(r"https?://(www\.)?linkedin\.com/[a-zA-Z0-9-_/]+|www\.linkedin\.com/[a-zA-Z0-9-_/]+", text),
        'Facebook': re.findall(r'https?://(?:www\.)?facebook\.com/[^\s]+', text),
        'X': re.findall(r'https?://(?:www\.)?twitter\.com/[^\s]+', text),
        'GitHub':re.findall(r"https?://github\.com/[a-zA-Z0-9_-]+",text)
    }
    return profiles

def extract_education(text):
    education = []
    
    # Pattern 1: Enhanced table format parsing (Lakshman Gupta's DOCX)
    table_pattern = r'\b(Pursuing BCA|10\+2|High School)\b\s*\|\s*(\d+%)\s*\|\s*.*?(\d{4})\s*\|\s*(.*?)(?=\n\||\n\+)'
    table_matches = re.finditer(table_pattern, text)
    for match in table_matches:
        education.append({
            'Degree': match.group(1).strip(),
            'Institution': match.group(4).strip(),
            'Year': match.group(3).strip()  # Now correctly captures year through "> 2022"
        })
    
    # Pattern 2: Enhanced qualification parsing (MD Gufran's PDF)
    if not education:
        edu_match = re.search(
            r'Educational Qualifications\s*:\s*(.*?)\s*,\s*(.*?)\s*[–-]\s*(\d{4}\s*[–-]\s*\d{4})',
            text
        )
        if edu_match:
            education.append({
                'Degree': edu_match.group(1).strip(),
                'Institution': edu_match.group(2).strip(),
                'Year': edu_match.group(3).replace('–', '-').strip()
            })

    return education or [{
        "Degree": "Not provided",
        "Institution": "Not provided",
        "Year": "Not provided"
    }]


def extract_experience(text):
    experience = []
    # Improved regex to handle different date separators and company formats
    exp_pattern = r'(?i)(Internship|Experience|Work History)[\s:-]*([^\n—–-]+)[—–-]\s*([^\n(]+?)\s*\((.*?)\)'
    exp_matches = re.finditer(exp_pattern, text)
    
    for match in exp_matches:
        experience.append({
            'Job Title': match.group(2).strip(),
            'Company': match.group(3).strip(),
            'Duration': match.group(4).strip(),
            'Description': extract_bullet_points(text)
        })

    # Fallback for project-based experience
    if not experience:
        project_match = re.search(
            r'Project\s*:\s*(.+?)\s*Environment\s*:\s*(.+?)\s*Project Description\s*:\s*(.+?)(?=\n\w+)', 
            text, re.DOTALL
        )
        if project_match:
            experience.append({
                'Job Title': f"Project: {project_match.group(1).strip()}",
                'Company': "Personal Project",
                'Duration': "Not Specified",
                'Description': [project_match.group(3).strip()]
            })

    return experience[:1] or [{
        "Job Title": "Not provided",
        "Company": "Not provided",
        "Duration": "Not provided",
        "Description": []
    }]

def extract_bullet_points(text):
    # Capture multi-line bullet points
    return re.findall(r'(?:•|\-|)\s*(.+?)(?=\n\s*(?:•|\-||\w+))', text, re.DOTALL)

def process_resume(file_path):
    # Extract text from the resume based on its type
    if file_path.endswith('.pdf'):
        text = extract_text_from_pdf(file_path)
    elif file_path.endswith('.docx'):
        text = extract_text_from_docx(file_path)
    elif file_path.endswith('.txt'):
        text = extract_text_from_txt(file_path)
    else:
        raise ValueError("Unsupported file format")

    print("\n====== Extracted Resume Text ======\n")
    print(text)
    print("\n====================================\n")

    dob = extract_dob(text)
   
    # Extract personal details
    names = extract_name(text)
    email = extract_email(text)
    phone = extract_phone_number(text)
    profiles = extract_profiles(text)
    experience = extract_experience(text)
    education = extract_education(text)
    address= extract_address(text)
    dob = extract_dob(text)
    skills = extract_skills(text)
    father_name = extract_father_name(text)
    gender = extract_gender(text)
    languages = extract_languages(text)
    certifications = extract_certifications(text)

    # Handle empty lists safely
    def get_first_or_default(lst, default="Not provided"):
        return lst[0] if lst else default

    # Structure the information
    structured_info = {
        "Skills":skills,
        "Certifications": certifications,
        "Personal Details": {
            "First Name": names[0] if names else "Not provided",
            "Last Name": names[1] if len(names) > 1 else "Not provided",
            "Email": email[0] if email else "Not provided",
            "Phone Number": phone[0] if phone else "Not provided",
            "Address": address if address else "Not provided",
            "Date Of Birth" : dob if dob else "Not provided",
            "Father Name": father_name,
            "Gender": gender,
            "Languages": languages,
        },
         "Experience": [
        {
            "Job Title": exp.get('Job Title', 'Not provided'),
            "Company": exp.get('Company', 'Not provided'),
            "From date": exp.get('Duration', '').split('-')[0].strip() if '-' in exp.get('Duration', '') else exp.get('Duration', 'Not provided'),
            "To date": exp.get('Duration', '').split('-')[-1].strip() if '-' in exp.get('Duration', '') else "Present",
            "Job Description": ' | '.join(exp.get('Description', []))
        }
        for exp in experience
    ],
        "Education": [
        {
            "Institution": edu.get('Institution', 'Not provided'),
            "Degree": edu.get('Degree', 'Not provided'),
            "From": edu.get('Year', '').split('-')[0].strip() if '-' in edu.get('Year', '') else "Not provided",
            "To": edu.get('Year', '').split('-')[-1].strip() if '-' in edu.get('Year', '') else "Not provided"
        }
        for edu in education
    ],
        "Your Profiles": {
            "LinkedIn": get_first_or_default(profiles.get("LinkedIn")),
            "GitHub": get_first_or_default(profiles.get("GitHub")),
            "Facebook": get_first_or_default(profiles.get("Facebook")),
            "X (fka Twitter)": get_first_or_default(profiles.get("X")),
            "Website": get_first_or_default(profiles.get("Website"))
        }
    }

    return structured_info

def save_to_excel(data, file_name="output.xlsx"):
    # Create DataFrame with full column structure
    df = pd.DataFrame(columns=[
        'First Name', 'Middle Name', 'Last Name', 'Father Name', 'DOB', 'Gender',
        'Email ID', 'Languages', 'Residential Address', 'Current Address', 'Phone No',
        'Total Experience', 'Job Title', 'Company', 'Office Location', 'Job Description',
        'From date', 'To date',  # Experience 1
        'Institution', 'Major', 'Degree', 'School Location', 'Description', 'From date', 'To date',  # Education 1
        'Certification 1', 'Certification 2', 'Certification 3', 'Key skills'
    ])

    # Populate personal details
    personal = data['Personal Details']
    df.at[0, 'First Name'] = personal.get('First Name', '')
    df.at[0, 'Last Name'] = personal.get('Last Name', '')
    df.at[0, 'Email ID'] = personal.get('Email', '').replace('-', '') 
    df.at[0, 'Phone No'] = personal.get('Phone Number', '')
    df.at[0, 'DOB'] = personal.get('Date Of Birth', '')
    df.at[0, 'Residential Address'] = personal.get('Address', '')
    df.at[0, 'Current Address'] = personal.get('Address', '')
    df.at[0, 'Father Name'] = personal.get('Father Name', '')
    df.at[0, 'Gender'] = personal.get('Gender', '')
    df.at[0, 'Languages'] = ', '.join(personal.get('Languages', []))

    # Populate Experience
    if data.get('Experience'):
        for i, exp in enumerate(data['Experience'][:1]):  # First experience only
            df.at[0, 'Job Title'] = exp.get('Job Title', '')
            df.at[0, 'Company'] = exp.get('Company', '')
            df.at[0, 'Job Description'] = exp.get('Job Description', '')
            df.at[0, 'From date'] = exp.get('From date', '')
            df.at[0, 'To date'] = exp.get('To date', '')

    # Populate Education
    if data.get('Education'):
        for i, edu in enumerate(data['Education'][:1]):  # First education only
            df.at[0, 'Institution'] = edu.get('Institution', '')
            df.at[0, 'Degree'] = edu.get('Degree', '')
            df.at[0, 'From date'] = edu.get('From', '')
            df.at[0, 'To date'] = edu.get('To', '')

    # Populate Certifications (Add this new section)
    if 'Certifications' in data:
        for i, cert in enumerate(data['Certifications'][:3]):
            df.at[0, f'Certification {i+1}'] = cert

    # Populate Skills
    df.at[0, 'Key skills'] = ', '.join(data.get('Skills', []))

    df.to_excel(file_name, index=False, sheet_name='Sheet1')
   
@app.route('/', methods=['GET', 'POST'])
def home():
    if request.method == 'POST':
        if 'file' not in request.files:
            return "No file part", 400

        file = request.files['file']

        if file and allowed_file(file.filename):
            filename = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
            file.save(filename)

            # Process the file
            structured_info = process_resume(filename)

            save_to_excel(structured_info, file_name="extracted_resume_data.xlsx")

            # Display the extracted information
            return {
                "message": "Data extracted and saved to extracted_resume_data.xlsx",
                "extracted_info": structured_info
            }

    # Display the file upload form
    return """
    <!DOCTYPE html>
    <html>
    <head>
        <title>Resume Extraction API</title>
    </head>
    <body>
        <h1>Upload a Resume</h1>
        <form action="/" method="POST" enctype="multipart/form-data">
            <input type="file" name="file" required>
            <button type="submit">Upload and Extract</button>
        </form>
    </body>
    </html>
    """
if __name__ == '__main__':
    app.run(debug=True)